# -*- coding: utf-8 -*-
"""
Created on Fri Aug  9 16:22:23 2024

@author: dbk
"""
import docx
from docx import Document
from docx.shared import RGBColor, Inches
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
 
# 创建一个Document对象

doc = docx.Document() 
doc.save("example.docx")

doc = docx.Document("example.docx")

# 创让一个新的段落
paragraph = doc.add_paragraph()
# 解加女本到聚落
run = paragraph.add_run("点击这里访问我的博客")
# 将段落转换为超链技
hyperlink = docx.oxml.shared.OxmlElement("w:hyperlink")

hyperlink.set(docx.oxml.ns.qn("w:anchor"),"FF8822")
# 将超链该添加测段落中
run._r.append(hyperlink)
# 保存父档
doc.save("example.docx")
