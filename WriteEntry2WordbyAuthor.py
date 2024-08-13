# -*- coding: utf-8 -*-
"""
Created on Mon Sep 02 10:15:07 2019

@author: SJB-083

#适用于数据中无样式的情况
"""

import re
from docx import Document
import sys
sys.path.append('G:\ECPH_LY\MyPythonProject(python3)')
import Head
def GetTagContent(regExp,string):

    content = Head.re.findall(regExp,string)
    if len(content) > 0:
        return content[0]
    else:
        return ''
        
if __name__ == "__main__":
    
    data_dir = u'G:/ECPH_LY/Data/协助同事/王瑜/唐诗PDF转WORD（20210823）'
    file_list = Head.GetFileList(data_dir,['.xml'])
    entry_f = r'<entry>[\s\S]*?</entry>'
    author_f = r'<AUTHOR>(.*?)</AUTHOR>'
    itemcn_f = r'<itemcn>(.*?)</itemcn>'
    body_f = r'<body>([\s\S]*?)</body>'
    entry_r = re.compile(entry_f)
    author_r = re.compile(author_f)
    itemcn_r = re.compile(itemcn_f)
    body_r = re.compile(body_f)

    for xml_file in file_list:
        xml_data = Head.ReadFile(xml_file)
        xml_data = xml_data.replace('<p>','')
        xml_data = xml_data.replace('</p>','')
        entry_list = re.findall(entry_r,xml_data)
        document = Document()
        for entry in entry_list:
            itemcn = GetTagContent(itemcn_r,entry)  
            body = GetTagContent(body_r,entry)
            author = GetTagContent(author_r,entry)
            document.add_heading(itemcn, level=2)
            document.add_paragraph(body)
            document.add_paragraph(author)
        document.save(xml_file.replace('xml','.docx'))

