# -*- coding: utf-8 -*-
"""
Created on Tue Apr 07 15:21:43 2020

@author: Liang Yan

#适用于数据中无样式无图片的情况
"""
import re
from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor
import sys
sys.path.append('G:\ECPH_LY\MyPythonProject(python3)')
import Head
def GetTagContent(regExp,string):

    content = Head.re.findall(regExp,string)
    if len(content) > 0:
        return content[0]
    else:
        return ''  
       
       
if __name__ == "__main__":
    
    data_dir = 'G:/ECPH_LY/Data/协助同事/！三版内容中心/抽取送审人物/test/est'
    file_list = Head.GetFileList(data_dir,['.xml'])
    entry_f = r'<entry>[\s\S]*?</entry>'
    itemcn_f = r'<itemcn>(.*?)</itemcn>'
    itemen_f = r'<itemen>(.*?)</itemen>'
    itempy_f = r'<itempy>(.*?)</itempy>'
    body_f = r'<body>\n([\s\S]*?)</body>'
    img_f = r'<IMG src=.*?></IMG>'
    #img_file_f = r'<IMG src="(.*?)"></IMG>'
    entry_r = re.compile(entry_f)   
    itemcn_r = re.compile(itemcn_f)
    itemen_r = re.compile(itemen_f)
    itempy_r = re.compile(itempy_f)
    body_r = re.compile(body_f)
    img_r = re.compile(img_f)
    personAge_r = re.compile(r'<PersonAge>(.*?)</PersonAge>')
    #img_file_r = re.compile(img_file_f)    
    for xml_file in file_list:
        
        xml_data = Head.ReadFile(xml_file)
        xml_data = xml_data.replace('<p>','')
        xml_data = xml_data.replace('</p>','')
        #得到xml图片链接列表 <IMG src="1.png"></IMG>，用于在word里替换成{{img1}}
        img_src_list = re.findall(img_r,xml_data)
            
        entry_list = re.findall(entry_r,xml_data)
        document = Document()
        document.styles['Normal'].font.name=u'宋体'
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        '''
        for entry in entry_list:
            #处理每一个条目
            itemcn = GetTagContent(itemcn_r,entry)  
            itemen = GetTagContent(itemen_r,entry)
            body = GetTagContent(body_r,entry)
            document.add_heading(itemcn, level=2)
            if itemen != '':
                document.add_heading(itemen, level=2)            
            document.add_paragraph(body)
        '''
        for entry in entry_list:
            #处理每一个条目
            itemcn = GetTagContent(itemcn_r,entry)  
            itempy = GetTagContent(itempy_r,entry)
            personAge = GetTagContent(personAge_r,entry)
            itemen = GetTagContent(itemen_r,entry)
            body = GetTagContent(body_r,entry)
            document.add_paragraph(itemcn)    
            document.add_paragraph(personAge) 
            #document.add_paragraph(itempy)    
            #document.add_paragraph(itemen) 
            document.add_paragraph(body)
      

        #存储为中间word文件
        result_doc = xml_file.replace('.xml','.docx')
        document.save(result_doc)

    


