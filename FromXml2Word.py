# -*- coding: utf-8 -*-
"""
Created on Thu Sep  9 14:51:46 2021

@author: dbk
"""
import zipfile
import os
import sys
sys.path.append('G:\ECPH_LY\MyPythonProject(python3)')
import Head
import re
from docx import Document
#from docxtpl import DocxTemplate, InlineImage
from docx.oxml.ns import qn
import jinja2
from docx.shared import Pt
from docx.shared import RGBColor, Inches

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT#导入段落对齐包
from docx.shared import Cm
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.text import WD_ALIGN_PARAGRAPH #导入反落对产
from docx.enum.table import WD_ALIGN_VERTICAL#导入单元格垂直对齐
from docx.oxml import OxmlElement,ns
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml



def extract(path, tempdir):
    obj = zipfile.ZipFile(path)
    obj.extractall(tempdir)

def GetTagContent(regExp,string):
    
    content = Head.re.findall(regExp,string)
    if len(content) == 1:
        string = content[0]

        return string
    else:
        return ''
    
def wLine(document,text,ls=1.5,ft=12):

    para = document.add_paragraph()
    para.paragraph_format.line_spacing = ls
    inUpLo=False
    sub,sup,xie=False,False,False
    i=0
    while i<len(text):
        ti=text[i]      
        if ti in ['，','d']:
            ftn ='宋体'
        else:
            ftn = 'Times New Roman'
        if text[i:i+6].startswith('<SUP>'):
            sup,inUpLo=True,True
            i+=5
            continue
        elif text[i:i+6].startswith('<SUB>'):
            sub,inUpLo=True,True
            i+=5
            continue
        elif text[i:i+4].startswith('<I>'):
            xie,inUpLo=True,True
            i+=3
            continue
        else:
            if text[i:i+7].startswith('</SUP>') and inUpLo:
                inUpLo=False
                i+=6
                sub,sup,xie=False,False,False
                continue
            elif text[i:i+7].startswith('</SUB>') and inUpLo:
                inUpLo=False
                i+=6
                sub,sup,xie=False,False,False
                continue
            elif text[i:i+5].startswith('</I>') and inUpLo:
                inUpLo=False
                i+=4
                sub,sup,xie=False,False,False
                continue
            
        r = para.add_run(ti)
        r.font.superscript = sup
        r.font.subscript = sub
        r.font.italic = xie
        r.font.name = ftn
        r.font.size = Pt(ft)
        para.alignment = WD_TAB_ALIGNMENT.LEFT
        r.font.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
        i = i + 1
#----------------------------------------------------------------------
if __name__ == "__main__":
    
    '''
    #处理单选题
    data_dir = 'G:/ECPH_LY/Data/项目相关/4中国科技馆题库/word稿件转xml后new/抽出来的题' 
    xml_dir = data_dir + '/xml'
    word_dir = data_dir + '/word2'

    item_r = re.compile(r'<item>([\s\S]*?)</item>')
    quesion_r = r'<p><question>([\s\S]*?)</question></p>'
    option_r = r'<options>([\s\S]*?)</options>'
    answer_r = r'<p><answer>([\s\S]*?)</answer></p>'
    analysis_r = r'<p><analysis>([\s\S]*?)</analysis></p>'
    subject_r = r'<p><subject>(.*?)</subject></p>'
    category_r = r'<p><category>(.*?)</category></p>'
    difficulty_level_r = r'<p><difficulty_level>(.*?)</difficulty_level></p>'
    gradation_r = r'<p><gradation>(.*?)</gradation></p>'
    knowledge_extension_r = r'<p><knowledge_extension>([\s\S]*?)</knowledge_extension></p>'
    
    
    xml_file_list = Head.GetFileList( xml_dir, ['.xml'])
    for xml_f in xml_file_list:
        xml_data = Head.ReadFile(xml_f)
        result_doc = xml_f.replace('.xml','.docx')
        result_doc = result_doc.replace(xml_dir,word_dir)
        document = Document()
        #设置中文字体
        microsoft_font = '黑体'
        area = qn('w:eastAsia')
        document.styles['Normal'].font.name = microsoft_font
        document.styles['Normal']._element.rPr.rFonts.set(area, microsoft_font)
        document.styles['Normal'].font.size = Pt(12)
        item_list = re.findall(item_r,xml_data)
        for i in range(len(item_list)):
            item = item_list[i]
            question = GetTagContent(quesion_r,item)
            option = GetTagContent(option_r,item)
            answer = GetTagContent(answer_r,item)
            analysis = GetTagContent(analysis_r,item)
            subject = GetTagContent(subject_r,item)
            category = GetTagContent(category_r,item)
            gradation = GetTagContent(gradation_r,item)
            difficulty_level = GetTagContent(difficulty_level_r,item)
            knowledge_extension = GetTagContent(knowledge_extension_r,item)
            document.add_paragraph(str(i+1) + '.【题干】'+question)
            document.add_paragraph('【选项】'+option)
            document.add_paragraph('【答案】'+answer)
            document.add_paragraph('【解析】'+analysis)
            document.add_paragraph('【学科分类】'+subject)
            document.add_paragraph('【科学分类】'+category)
            document.add_paragraph('【难易程度】'+difficulty_level)
            document.add_paragraph('【阶段分类】'+gradation)
            #document.add_paragraph('【知识延伸】'+knowledge_extension)
            #让链接能点击
            knowledge_extension_list = knowledge_extension.split('\n')
            for k in range(len(knowledge_extension_list)):
                knowledge = knowledge_extension_list[k]
                if k == 0:
                    document.add_paragraph('【知识延伸】'+knowledge)
                else:
                    if knowledge[:4] == 'http':
                        p = document.add_paragraph()
                        url_run=p.add_run(knowledge)
                        url_run.font.color.rgb = RGBColor(0, 0, 255)
                        url_run.hyperlink = knowledge
                    else:
                        document.add_paragraph(knowledge)   
                        
        document.save(result_doc)
    '''
    '''
    #处理材料题
    data_dir = 'G:/ECPH_LY/Data/项目相关/4中国科技馆题库/加id后/材料题/打乱顺序合并' 
    xml_dir = data_dir
    word_dir = data_dir + '/word'

    item_r = re.compile(r'<item>([\s\S]*?)</item>')
    desc_r = r'<p><description>([\s\S]*?)</description></p>'
    quesion_r = r'<p><question>([\s\S]*?)</question></p>'
    #option_r = r'<options>([\s\S]*?)</options>'
    answer_r = r'<p><answer>([\s\S]*?)</answer></p>'
    analysis_r = r'<p><analysis>([\s\S]*?)</analysis></p>'
    subject_r = r'<p><subject>(.*?)</subject></p>'
    category_r = r'<p><category>(.*?)</category></p>'
    difficulty_level_r = r'<p><difficulty_level>(.*?)</difficulty_level></p>'
    gradation_r = r'<p><gradation>(.*?)</gradation></p>'
    knowledge_extension_r = r'<p><knowledge_extension>([\s\S]*?)</knowledge_extension></p>'
    
    
    xml_file_list = Head.GetFileList( xml_dir, ['.xml'])
    for xml_f in xml_file_list:
        xml_data = Head.ReadFile(xml_f)
        result_doc = xml_f.replace('.xml','.docx')
        result_doc = result_doc.replace(xml_dir,word_dir)
        document = Document()
        #设置中文字体
        microsoft_font = '黑体'
        area = qn('w:eastAsia')
        document.styles['Normal'].font.name = microsoft_font
        document.styles['Normal']._element.rPr.rFonts.set(area, microsoft_font)
        document.styles['Normal'].font.size = Pt(12)
        item_list = re.findall(item_r,xml_data)
        for i in range(len(item_list)):
            item = item_list[i]
            description = GetTagContent(desc_r,item)
            question = GetTagContent(quesion_r,item)
            #option = GetTagContent(option_r,item)
            answer = GetTagContent(answer_r,item)
            analysis = GetTagContent(analysis_r,item)
            subject = GetTagContent(subject_r,item)
            category = GetTagContent(category_r,item)
            gradation = GetTagContent(gradation_r,item)
            difficulty_level = GetTagContent(difficulty_level_r,item)
            knowledge_extension = GetTagContent(knowledge_extension_r,item)
            document.add_paragraph(str(i+1) + '.【题干】'+description)
            document.add_paragraph('【问题】'+question)
            #document.add_paragraph('【选项】'+option)
            document.add_paragraph('【答案】'+answer)
            document.add_paragraph('【解析】'+analysis)
            document.add_paragraph('【学科分类】'+subject)
            document.add_paragraph('【科学分类】'+category)
            #document.add_paragraph('【难易程度】'+difficulty_level)
            #document.add_paragraph('【阶段分类】'+gradation)
            document.add_paragraph('【知识延伸】'+knowledge_extension)

        document.save(result_doc)
        '''
    #处理单选题包含上下标斜体
    data_dir = 'G:/ECPH_LY/Data/项目相关/4中国科技馆题库/word稿件转xml后new/抽出来的题' 
    xml_dir = data_dir + '/xml'
    word_dir = data_dir + '/word3'

    item_r = re.compile(r'<item>([\s\S]*?)</item>')
    quesion_r = r'<p><question>([\s\S]*?)</question></p>'
    option_r = r'<options>([\s\S]*?)</options>'
    answer_r = r'<p><answer>([\s\S]*?)</answer></p>'
    analysis_r = r'<p><analysis>([\s\S]*?)</analysis></p>'
    subject_r = r'<p><subject>(.*?)</subject></p>'
    category_r = r'<p><category>(.*?)</category></p>'
    difficulty_level_r = r'<p><difficulty_level>(.*?)</difficulty_level></p>'
    gradation_r = r'<p><gradation>(.*?)</gradation></p>'
    knowledge_extension_r = r'<p><knowledge_extension>([\s\S]*?)</knowledge_extension></p>'
    
    
    xml_file_list = Head.GetFileList( xml_dir, ['.xml'])
    for xml_f in xml_file_list:
        xml_data = Head.ReadFile(xml_f)
        result_doc = xml_f.replace('.xml','.docx')
        result_doc = result_doc.replace(xml_dir,word_dir)
        document = Document()
        #设置中文字体
        microsoft_font = '黑体'
        area = qn('w:eastAsia')
        document.styles['Normal'].font.name = microsoft_font
        document.styles['Normal']._element.rPr.rFonts.set(area, microsoft_font)
        document.styles['Normal'].font.size = Pt(12)
        item_list = re.findall(item_r,xml_data)
        for i in range(len(item_list)):
            item = item_list[i]
            question = GetTagContent(quesion_r,item)
            option = GetTagContent(option_r,item)
            answer = GetTagContent(answer_r,item)
            analysis = GetTagContent(analysis_r,item)
            subject = GetTagContent(subject_r,item)
            category = GetTagContent(category_r,item)
            gradation = GetTagContent(gradation_r,item)
            difficulty_level = GetTagContent(difficulty_level_r,item)
            knowledge_extension = GetTagContent(knowledge_extension_r,item)
            wLine(document,str(i+1) + '.【题干】'+question)
            wLine(document,'【选项】'+option)
            wLine(document,'【答案】'+answer)
            wLine(document,'【解析】'+analysis)
            wLine(document,'【学科分类】'+subject)
            wLine(document,'【科学分类】'+category)
            wLine(document,'【难易程度】'+difficulty_level)
            wLine(document,'【阶段分类】'+gradation)

            #document.add_paragraph('【知识延伸】'+knowledge_extension)
            #让链接能点击
            knowledge_extension_list = knowledge_extension.split('\n')
            for k in range(len(knowledge_extension_list)):
                knowledge = knowledge_extension_list[k]
                if k == 0:
                    wLine(document,'【知识延伸】'+knowledge)
                else:
                    if knowledge[:4] == 'http':
                        p = document.add_paragraph()
                        url_run=p.add_run(knowledge)
                        url_run.font.color.rgb = RGBColor(0, 0, 255)
                        url_run.hyperlink = knowledge
                    else:
                        wLine(document,knowledge)   
                        
        document.save(result_doc)
    
        
    