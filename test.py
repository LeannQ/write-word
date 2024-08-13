# -*- coding: utf-8 -*-
"""
Created on Fri Aug  9 10:28:59 2024

@author: dbk
"""

import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT#导入段落对齐包
from docx.shared import RGBColor # 设置字体的颜色
from docx.shared import Cm, Inches, Pt
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.text import WD_ALIGN_PARAGRAPH #导入反落对产
from docx.enum.table import WD_ALIGN_VERTICAL#导入单元格垂直对齐
from docx.oxml import OxmlElement,ns
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx import Document
from docx.shared import RGBColor, Pt
from docx.shared import Inches
from docx.oxml.ns import nsdecls
from docx.text.hyperlink import Hyperlink

'''
def wLine(document,text,ls=1.5,ft=12):

    para = document.add_paragraph()
    para.paragraph_format.line_spacing = ls
    inUpLo=False
    sub,sup,xie=False,False,False
    i=0
    while i<len(text):
        ti=text[i]      
        if ti in ['，','d']:
            ftn ='宋体'
        else:
            ftn = 'Times New Roman'
        if text[i:i+5].startswith('$up{'):
            sup,inUpLo=True,True
            i+=4
            continue
        elif text[i:i+5].startswith('$lo{'):
            sub,inUpLo=True,True
            i+=4
            continue
        elif text[i:i+5].startswith('$xe{'):
            xie,inUpLo=True,True
            i+=4
            continue
        else:
            if text[i]=='}' and inUpLo:
                inUpLo=False
                i+=1
                sub,sup,xie=False,False,False
                continue
            
        r = para.add_run(ti)
        r.font.superscript = sup
        r.font.subscript = sub
        r.font.italic = xie
        r.font.name = ftn
        r.font.size = Pt(ft)
        para.alignment = WD_TAB_ALIGNMENT.LEFT
        r.font.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
        i = i + 1
'''
'''
def wLine(document,text,ls=1.5,ft=12):

    para = document.add_paragraph()
    para.paragraph_format.line_spacing = ls
    inUpLo=False
    sub,sup,xie=False,False,False
    i=0
    while i<len(text):
        ti=text[i]      
        if ti in ['，','d']:
            ftn ='宋体'
        else:
            ftn = 'Times New Roman'
        if text[i:i+6].startswith('<SUP>'):
            sup,inUpLo=True,True
            i+=5
            continue
        elif text[i:i+6].startswith('<SUB>'):
            sub,inUpLo=True,True
            i+=5
            continue
        elif text[i:i+4].startswith('<I>'):
            xie,inUpLo=True,True
            i+=3
            continue
        else:
            if text[i:i+7].startswith('</SUP>') and inUpLo:
                inUpLo=False
                i+=6
                sub,sup,xie=False,False,False
                continue
            elif text[i:i+7].startswith('</SUB>') and inUpLo:
                inUpLo=False
                i+=6
                sub,sup,xie=False,False,False
                continue
            elif text[i:i+5].startswith('</I>') and inUpLo:
                inUpLo=False
                i+=4
                sub,sup,xie=False,False,False
                continue
            
        r = para.add_run(ti)
        r.font.superscript = sup
        r.font.subscript = sub
        r.font.italic = xie
        r.font.name = ftn
        r.font.size = Pt(ft)
        para.alignment = WD_TAB_ALIGNMENT.LEFT
        r.font.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
        i = i + 1

if __name__=='__main__':
    
    doc = docx.Document()
    wLine(doc,'这是一个测试“斜体<I>对</I>，没错，ye<SUB>s</SUB>，sir。”，这有个<SUP>上标</SUP>还有个标。\n或者你说呢<SUB>e</SUB>.')
    url = 'https://www.example.com'
    paragraph = doc.add_paragraph()
    hyperlink = Hyperlink(paragraph, url)
 
    # 添加超链接的文本
    hyperlink.runs(url).font.color.rgb = RGBColor(0, 0, 255)

    # 保存文档
    doc.save('test.docx')
'''    
from docx.enum.text import WD_UNDERLINE
def add_hyperlink(paragraph, url, text, color, underline):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
      c = docx.oxml.shared.OxmlElement('w:color')
      c.set(docx.oxml.shared.qn('w:val'), color)
      rPr.append(c)

    # Remove underlining if it is requested
    if not underline:
      u = docx.oxml.shared.OxmlElement('w:u')
      u.set(docx.oxml.shared.qn('w:val'), 'none')
      rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    #underline_style = parse_xml(r'<u w:val="single"/>')
    #hyperlink.font.underline = underline_style
    paragraph._p.append(hyperlink)

    return hyperlink


document = docx.Document()
p = document.add_paragraph()

#add a hyperlink with the normal formatting (blue underline)
hyperlink = add_hyperlink(p, 'http://www.google.com', 'http://www.google.com', None, False)

#add a hyperlink with a custom color and no underline
hyperlink = add_hyperlink(p, 'http://www.google.com', 'Google', 'FF8822', False)

document.save('demo.docx')