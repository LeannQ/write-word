# -*- coding: utf-8 -*-
"""
Created on Fri Sep 06 10:41:47 2019

@author: SJB-083
"""

# -*- coding: utf-8 -*-
"""
Created on Mon Sep 02 10:15:07 2019

@author: SJB-083

#适用于数据中有图片无样式的情况
"""
import re
from docx import Document
from docxtpl import DocxTemplate, InlineImage
import jinja2
import sys
sys.path.append('G:\ECPH_LY\MyPythonProject(python3)')
import Head
def GetTagContent(regExp,string):

    content = Head.re.findall(regExp,string)
    if len(content) > 0:
        return content[0]
    else:
        return ''
        
if __name__ == "__main__":
    
    data_dir = u'G:/ECPH_LY/Data/协助同事/王瑜/唐诗PDF转WORD（20210823）/拆分后'
    file_list = Head.GetFileList(data_dir,['.xml'])
    entry_f = r'<entry>[\s\S]*?</entry>'
    itemcn_f = r'<itemcn>(.*?)</itemcn>'
    body_f = r'<body>([\s\S]*?)</body>'
    img_f = r'<IMG src=.*?></IMG>'
    img_file_f = r'<IMG src="(.*?)"></IMG>'
    entry_r = re.compile(entry_f)   
    itemcn_r = re.compile(itemcn_f)
    body_r = re.compile(body_f)
    img_r = re.compile(img_f)
    img_file_r = re.compile(img_file_f)
    
    for xml_file in file_list:
        
        xml_data = Head.ReadFile(xml_file)
        xml_data = xml_data.replace('<p>','')
        xml_data = xml_data.replace('</p>','')
        #得到xml图片链接列表 <IMG src="1.png"></IMG>，用于在word里替换成{{img1}}
        img_src_list = re.findall(img_r,xml_data)
        #得到图片文件名列表 1.png
        img_list = re.findall(img_file_r,xml_data)
        #将图片文件名调整至当前存储图片的路径 E:/ECPH_LY/Data/协助同事/王瑜/宋词/按作者拆分后/test/img/1.png
        img_list = [data_dir+'/img/'+x for x in img_list]

        #替换xml中的图片路径
        for i in range(len(img_list)):
            xml_data = xml_data.replace(img_src_list[i],'{{img'+str(i)+'}}')
            
        entry_list = re.findall(entry_r,xml_data)
        document = Document()
        for entry in entry_list:
            #处理每一个条目
            itemcn = GetTagContent(itemcn_r,entry)  
            body = GetTagContent(body_r,entry)
            document.add_heading(itemcn, level=2)
            document.add_paragraph(body)
        #存储为中间word文件
        result_doc = xml_file.replace('.xml','.docx')
        document.save(result_doc)
        #按照中间word文件模板替换
        tpl=DocxTemplate(result_doc)
        #创建dict
        key_list = []
        for i in range(len(img_list)):
            key_list.append('img'+str(i))
        img_list_new = [InlineImage(tpl,x) for x in img_list]
        #ccc = dict(zip(key_list,img_list))
        context = dict(zip(key_list,img_list_new))
    
        jinja_env = jinja2.Environment(autoescape=True)
        tpl.render(context, jinja_env)
        tpl.save(result_doc)
    


