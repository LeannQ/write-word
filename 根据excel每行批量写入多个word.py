# -*- coding: utf-8 -*-
"""
Created on Thu Sep 29 14:23:04 2022

@author: dbk
"""

# -*- coding: utf-8 -*-
"""
Created on Tue Apr 07 15:21:43 2020

@author: Liang Yan

#适用于数据中无样式无图片的情况
"""
import re
from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.shared import Inches,Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys
sys.path.append('G:\ECPH_LY\MyPythonProject(python3)')
import Head
def GetTagContent(regExp,string):

    content = Head.re.findall(regExp,string)
    if len(content) > 0:
        return content[0]
    else:
        return ''  
       
       
if __name__ == "__main__":
    
    data_dir = 'G:/ECPH_LY/Data/协助同事/高斐斐/科技中国'
    excel_file = data_dir + '/科技中国word.xlsx'
    
    itemcn_list = Head.getDataFromExcel(excel_file,0,0,1)
    content_list = Head.getDataFromExcel(excel_file,0,1,1)
    
    for i in range(len(itemcn_list)):
        
        itemcn = itemcn_list[i]
        word_file = data_dir + '/' + itemcn + '.docx'
        content = content_list[i]
   
        document = Document()
        document.styles['Normal'].font.name=u'宋体'
        document.styles['Normal'].font.size= Pt(10.5)
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        para_list = content.split('\n')
        for p in para_list:
            p = p.strip()
            para = document.add_paragraph(p)   
            para.paragraph_format.first_line_indent = Pt(20)
            para.paragraph_format.line_spacing = 1.5 
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.space_after = Pt(0)
        para = document.add_paragraph('（来源：新华社）')   
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT    
        document.save(word_file)

    


